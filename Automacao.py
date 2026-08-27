import pandas as pd
from docx import Document
import os

documento = Document(
    r"C:\Users\ferna\Documents\Códigos\Automação\Contrato.docx"
)

dados = {
    "{{RG}}": "4",
    "{{RPA}}": "2",
    "{{NOME_PESSOA}}": "gozei gozei gozei",
    "{{MATRICULA_PESSOA}}": "121257,9",
    "{{HORAS_AULA}}": "195",
    "{{HORAS_AULA_EXTENSO}}": "zerocento e zero",
    "{{TURNO_PESSOA}}": "narde",
    "{{EMAIL_REGIONAL}}": "scpiniciaisregional7@educ.rec.br"
}


for tabela in documento.tables:
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragraph in celula.paragraphs:

                for condicional, valor in dados.items():

                    if condicional in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            condicional,
                            valor
                        )


documento.save(
    r"C:\Users\ferna\Documents\Códigos\Automação\Contrato_atualizado.docx"
)

print("Contrato atualizado com sucesso!")
    
